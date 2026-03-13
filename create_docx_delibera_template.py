#!/usr/bin/env python3
"""
Crea un template DOCX completo per la Delibera ASL.
Contiene:
1. Placeholder semplici ($$nome$$)
2. Placeholder complessi (P_testo_obj, I_testo_obj) in tabelle
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_docx_template(output_path='template/ASL_Template_Delibera.docx'):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    
    # Intestazione e Simple Placeholders
    p1 = doc.add_paragraph('DELIBERA N. $$numeroproposta$$')
    p2 = doc.add_paragraph('Data: $$dataproposta$$')
    
    title = doc.add_heading('$$oggetto$$', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sezione Istruttoria
    doc.add_heading('RELAZIONE ISTRUTTORIA', level=2)
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    table1.cell(0, 0).paragraphs[0].text = 'I_testo_obj'
    
    # Sezione Proposta
    doc.add_heading('PROPOSTA', level=2)
    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'
    table2.cell(0, 0).paragraphs[0].text = 'P_testo_obj'
    
    # Sezione Delibera
    doc.add_heading('DELIBERA', level=2)
    table3 = doc.add_table(rows=1, cols=1)
    table3.style = 'Table Grid'
    table3.cell(0, 0).paragraphs[0].text = 'P_testo_obj_2'
    
    # Firme / End
    doc.add_paragraph('')
    doc.add_paragraph('Estensore: $$EstensoreNome$$')
    doc.add_paragraph('Direttore Generale: $$DirGenNome$$')
    
    doc.save(output_path)
    print(f"[OK] Template DOCX creato: {output_path}")

if __name__ == '__main__':
    create_docx_template()
