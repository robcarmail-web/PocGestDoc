#!/usr/bin/env python3
"""
Inietta contenuto in template DOCX.
Approccio: copia XML direttamente dal sorgente al template usando lxml,
preservando liste, tabelle, formattazione completa.
"""
import zipfile
from io import BytesIO
from pathlib import Path
from copy import deepcopy
from lxml import etree
from typing import List, Optional

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = 'http://www.w3.org/XML/1998/namespace'


def _w(tag):
    return '{%s}%s' % (W, tag)


class DocxInjector:
    """Sostituisce P_testo_obj in un template DOCX con contenuto formattato."""

    def __init__(self, template_path: str):
        self.template_path = Path(template_path)

    # ------------------------------------------------------------------ #
    # Metodo principale: sorgente DOCX (upload)
    # ------------------------------------------------------------------ #

    def inject_from_docx(self, source_docx_path: str, placeholder: str = 'P_testo_obj') -> bytes:
        """
        Copia tutto il contenuto dal DOCX sorgente nel template,
        sostituendo il placeholder. Preserva liste, tabelle, formattazione.
        """
        # Leggi entrambi i file come ZIP
        template_files = self._read_zip(self.template_path)
        source_files = self._read_zip(source_docx_path)

        # Parsa i document.xml
        tpl_doc = etree.fromstring(template_files['word/document.xml'])
        src_doc = etree.fromstring(source_files['word/document.xml'])

        # Trova il body e il placeholder nel template
        tpl_body = tpl_doc.find('.//{%s}body' % W)
        placeholder_elem = self._find_placeholder(tpl_body, placeholder)

        if placeholder_elem is None:
            print(f'[WARNING] Placeholder "{placeholder}" not found')
            output = BytesIO()
            output.write(self._rebuild_zip(template_files))
            return output.getvalue()

        # Estrai tutti gli elementi dal body del sorgente (escluso sectPr finale)
        src_body = src_doc.find('.//{%s}body' % W)
        src_elements = [deepcopy(child) for child in src_body
                        if child.tag != _w('sectPr')]

        # Inserisci nel template nella posizione del placeholder
        for elem in reversed(src_elements):
            placeholder_elem.addnext(elem)

        # Rimuovi il placeholder
        tpl_body.remove(placeholder_elem)

        # Copia numbering.xml se presente nel sorgente
        if 'word/numbering.xml' in source_files:
            template_files['word/numbering.xml'] = source_files['word/numbering.xml']
            # Aggiorna il manifest
            template_files = self._update_manifest(template_files, 'word/numbering.xml',
                                                    'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml')
            # Aggiorna le relazioni
            template_files = self._update_rels(template_files, 'word/numbering.xml')

        # Copia stili del sorgente (merge base)
        if 'word/styles.xml' in source_files:
            template_files['word/styles.xml'] = self._merge_styles(
                template_files.get('word/styles.xml'),
                source_files['word/styles.xml']
            )

        # Aggiorna il document.xml nel template
        template_files['word/document.xml'] = etree.tostring(
            tpl_doc, xml_declaration=True, encoding='UTF-8', standalone=True
        )

        return self._rebuild_zip(template_files)

    # ------------------------------------------------------------------ #
    # Metodo alternativo: lista di paragrafi (editor web)
    # ------------------------------------------------------------------ #

    def inject_from_paragraphs(self, paragraph_list: List[dict], placeholder: str = 'P_testo_obj') -> bytes:
        """
        Costruisce elementi XML da paragraph_list e li inietta nel template.
        Usato per l'editor web (TipTap).
        """
        # Crea un DOCX temporaneo con i paragrafi usando python-docx
        # (usa il default template che include stili lista)
        temp_bytes = self._build_temp_docx(paragraph_list)

        # Poi inietta il contenuto di quel DOCX nel template
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(temp_bytes)
            temp_path = tmp.name
        try:
            return self.inject_from_docx(temp_path, placeholder)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def _build_temp_docx(self, paragraph_list: List[dict]) -> bytes:
        """Costruisce un DOCX temporaneo con i paragrafi usando python-docx."""
        from docx import Document
        doc = Document()

        for para_info in paragraph_list:
            if para_info['type'] == 'paragraph':
                para = doc.add_paragraph()
                try:
                    para.style = para_info.get('style', 'Normal')
                except Exception:
                    para.style = 'Normal'

                for run_info in para_info.get('runs', []):
                    run = para.add_run(run_info['text'])
                    if run_info.get('bold'):
                        run.bold = True
                    if run_info.get('italic'):
                        run.italic = True
                    if run_info.get('underline'):
                        run.underline = True

            elif para_info['type'] == 'list':
                list_type = para_info['list_type']
                for item in para_info.get('items', []):
                    para = doc.add_paragraph()
                    try:
                        para.style = 'List Bullet' if list_type == 'bullet' else 'List Number'
                    except Exception:
                        para.style = 'Normal'

                    for run_info in item.get('runs', []):
                        run = para.add_run(run_info['text'])
                        if run_info.get('bold'):
                            run.bold = True
                        if run_info.get('italic'):
                            run.italic = True
                        if run_info.get('underline'):
                            run.underline = True

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------ #
    # Utility
    # ------------------------------------------------------------------ #

    def _read_zip(self, path) -> dict:
        files = {}
        with zipfile.ZipFile(path, 'r') as zf:
            for name in zf.namelist():
                files[name] = zf.read(name)
        return files

    def _rebuild_zip(self, files: dict) -> bytes:
        output = BytesIO()
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
            # mimetype deve essere primo e non compresso
            if '[Content_Types].xml' in files:
                zf.writestr('[Content_Types].xml', files['[Content_Types].xml'])
            for name, data in files.items():
                if name != '[Content_Types].xml':
                    zf.writestr(name, data)
        return output.getvalue()

    def _find_placeholder(self, body, placeholder: str):
        """Cerca il paragrafo che contiene il placeholder."""
        for p in body.iter(_w('p')):
            text = ''.join(t.text or '' for t in p.iter(_w('t')))
            if placeholder in text:
                return p
        return None

    def _merge_styles(self, tpl_styles_bytes: Optional[bytes], src_styles_bytes: bytes) -> bytes:
        """Aggiunge stili del sorgente che non esistono nel template."""
        if not tpl_styles_bytes:
            return src_styles_bytes

        tpl_root = etree.fromstring(tpl_styles_bytes)
        src_root = etree.fromstring(src_styles_bytes)

        S = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # Raccogli styleId già presenti nel template
        existing_ids = set()
        for style in tpl_root.findall('.//{%s}style' % S):
            sid = style.get('{%s}styleId' % S)
            if sid:
                existing_ids.add(sid)

        # Aggiungi stili mancanti dal sorgente
        for style in src_root.findall('.//{%s}style' % S):
            sid = style.get('{%s}styleId' % S)
            if sid and sid not in existing_ids:
                tpl_root.append(deepcopy(style))

        return etree.tostring(tpl_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    def _update_manifest(self, files: dict, part_name: str, content_type: str) -> dict:
        """Aggiunge una entry al [Content_Types].xml se non esiste."""
        if '[Content_Types].xml' not in files:
            return files

        root = etree.fromstring(files['[Content_Types].xml'])
        CT = 'http://schemas.openxmlformats.org/package/2006/content-types'

        # Controlla se esiste già
        part_uri = '/' + part_name
        for override in root.findall('{%s}Override' % CT):
            if override.get('PartName') == part_uri:
                return files

        override = etree.SubElement(root, '{%s}Override' % CT)
        override.set('PartName', part_uri)
        override.set('ContentType', content_type)

        files['[Content_Types].xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        return files

    def _update_rels(self, files: dict, part_name: str) -> dict:
        """Aggiunge la relazione al word/_rels/document.xml.rels."""
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in files:
            return files

        root = etree.fromstring(files[rels_path])
        R = 'http://schemas.openxmlformats.org/package/2006/relationships'

        # Controlla se relazione esiste già
        target = part_name.replace('word/', '')
        for rel in root.findall('{%s}Relationship' % R):
            if rel.get('Target') == target:
                return files

        # Trova un ID disponibile
        existing_ids = {rel.get('Id') for rel in root.findall('{%s}Relationship' % R)}
        rid = 'rId' + str(len(existing_ids) + 20)

        rel = etree.SubElement(root, '{%s}Relationship' % R)
        rel.set('Id', rid)
        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
        rel.set('Target', target)

        files[rels_path] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        return files

    # ------------------------------------------------------------------ #
    # API pubblica: inject_content (routing automatico)
    # ------------------------------------------------------------------ #

    def inject_content(self, source, placeholder: str = 'P_testo_obj') -> bytes:
        """
        Router:
        - Se source è una stringa (path file), usa inject_from_docx
        - Se source è una lista di dict, usa inject_from_paragraphs
        """
        if isinstance(source, str):
            return self.inject_from_docx(source, placeholder)
        elif isinstance(source, list):
            return self.inject_from_paragraphs(source, placeholder)
        else:
            raise ValueError(f"Tipo sorgente non supportato: {type(source)}")

    def save(self, output_path: str, source, placeholder: str = 'P_testo_obj') -> None:
        """Salva il documento generato su disco."""
        docx_bytes = self.inject_content(source, placeholder)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'wb') as f:
            f.write(docx_bytes)
