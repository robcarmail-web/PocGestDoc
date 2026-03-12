"""
Converts DOCX document content to ODT XML fragments.
Extracts formatted text from python-docx paragraphs and converts to ODT.
"""
from typing import List, Dict, Tuple, Optional
from lxml import etree
from docx import Document

NS = {
    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
    'table': 'urn:oasis:names:tc:opendocument:xmlns:table:1.0',
    'style': 'urn:oasis:names:tc:opendocument:xmlns:style:1.0',
}

for prefix, uri in NS.items():
    etree.register_namespace(prefix, uri)


class DocxToODT:
    """Converts DOCX document to ODT XML fragments."""

    def __init__(self):
        self.styles_needed = set()

    def convert_file(self, docx_path: str) -> Tuple[List[str], str]:
        """
        Convert DOCX file to ODT fragments.

        Args:
            docx_path: Path to DOCX file

        Returns:
            (list of XML fragment strings, style definitions)
        """
        doc = Document(docx_path)
        fragments = []

        for para in doc.paragraphs:
            frag = self._convert_paragraph(para)
            if frag:
                fragments.append(frag)

        for table in doc.tables:
            frag = self._convert_table(table)
            if frag:
                fragments.append(frag)

        return fragments, self._generate_styles()

    def _convert_paragraph(self, para) -> Optional[str]:
        """Convert a DOCX paragraph to ODT XML."""
        odt_para = etree.Element('{' + NS['text'] + '}p')
        odt_para.set('{' + NS['text'] + '}style-name', 'Textbody')

        for run in para.runs:
            span = self._convert_run(run)
            if span is not None:
                odt_para.append(span)

        # Return XML string if has content
        if len(odt_para) > 0:
            return etree.tostring(odt_para, encoding='unicode')
        elif para.text.strip():
            # Paragraph with text but no explicit runs
            span = etree.Element('{' + NS['text'] + '}span')
            text_elem = etree.Element('{' + NS['text'] + '}t')
            text_elem.text = para.text
            span.append(text_elem)
            odt_para.append(span)
            return etree.tostring(odt_para, encoding='unicode')

        return None

    def _convert_run(self, run) -> Optional[etree._Element]:
        """Convert a DOCX run (text with formatting) to ODT span."""
        if not run.text:
            return None

        # Determine formatting
        is_bold = run.bold if run.bold is not None else False
        is_italic = run.italic if run.italic is not None else False
        is_underline = run.underline if run.underline is not None else False

        # Generate style name
        marks = []
        if is_bold:
            marks.append('bold')
        if is_italic:
            marks.append('italic')
        if is_underline:
            marks.append('underline')

        style_name = None
        if marks:
            style_name = 'T_' + '_'.join(sorted(marks))
            self.styles_needed.add(style_name)

        # Create span
        span = etree.Element('{' + NS['text'] + '}span')
        if style_name:
            span.set('{' + NS['text'] + '}style-name', style_name)

        text_elem = etree.Element('{' + NS['text'] + '}t')
        text_elem.text = run.text
        span.append(text_elem)

        return span

    def _convert_table(self, table) -> Optional[str]:
        """Convert DOCX table to ODT table."""
        odt_table = etree.Element('{' + NS['table'] + '}table')
        odt_table.set('{' + NS['table'] + '}name', 'Table1')
        odt_table.set('{' + NS['table'] + '}style-name', 'Table1')

        # Add column definitions
        if len(table.rows) > 0:
            num_cols = len(table.rows[0].cells)
            for _ in range(num_cols):
                col = etree.Element('{' + NS['table'] + '}table-column')
                col.set('{' + NS['table'] + '}style-name', 'Table1.A')
                odt_table.append(col)

        # Add rows
        for row in table.rows:
            odt_row = etree.Element('{' + NS['table'] + '}table-row')
            odt_row.set('{' + NS['table'] + '}style-name', 'Table1.1')

            for cell in row.cells:
                odt_cell = etree.Element('{' + NS['table'] + '}table-cell')
                odt_cell.set('{' + NS['table'] + '}style-name', 'Table1.A1')

                # Add cell content
                for para in cell.paragraphs:
                    para_xml = self._convert_paragraph(para)
                    if para_xml:
                        para_elem = etree.fromstring(para_xml)
                        odt_cell.append(para_elem)

                odt_row.append(odt_cell)

            odt_table.append(odt_row)

        return etree.tostring(odt_table, encoding='unicode')

    def _generate_styles(self) -> str:
        """Generate style definitions for automatic styles."""
        styles = []

        for style_name in sorted(self.styles_needed):
            parts = style_name[2:].split('_')  # Remove 'T_' prefix

            properties = []
            if 'bold' in parts:
                properties.append('fo:font-weight="bold"')
            if 'italic' in parts:
                properties.append('fo:font-style="italic"')
            if 'underline' in parts:
                properties.append('style:text-underline-style="solid"')

            props_str = ' '.join(properties)
            styles.append(f'<style:style style:name="{style_name}" style:family="text"><style:text-properties {props_str}/></style:style>')

        return '\n'.join(styles)
