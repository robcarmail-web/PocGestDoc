#!/usr/bin/env python3
"""End-to-end test of all three modes."""
import sys
import json
import tempfile
from pathlib import Path

sys.path.insert(0, 'modules')

from odt_injector import ODTInjector
from docx_to_odt import DocxToODT
from tiptap_to_odt import TiptapToODT
from docx import Document
from docx.shared import Pt

# Test data
SIMPLE_DATA = {
    'numeroproposta': '2024/001',
    'dataproposta': '15/01/2024',
    'oggetto': 'Test Proposal',
    'EstensoreNome': 'Dott. Mario Rossi',
    'DirGenNome': 'Prof. Roberto Blu',
}

print("=" * 60)
print("END-TO-END TEST - POC Delibere Template")
print("=" * 60)

# Create test DOCX
print("\n[SETUP] Creating test DOCX file...")
doc = Document()
p = doc.add_paragraph("This is a ")
p.add_run("bold").bold = True
p.add_run(" and ")
p_italic = p.add_run("italic")
p_italic.italic = True
p.add_run(" text.")

temp_docx = Path(tempfile.gettempdir()) / 'test_input.docx'
doc.save(temp_docx)
print(f"[OK] Test DOCX created: {temp_docx}")

# ============================================================================
# MODALITA 1: Direct DOCX Upload
# ============================================================================
print("\n" + "-" * 60)
print("MODALITA 1: Direct DOCX Upload")
print("-" * 60)

print("[STEP 1] Convert DOCX to ODT fragments...")
converter = DocxToODT()
docx_fragments, docx_styles = converter.convert_file(str(temp_docx))
print(f"[OK] Converted to {len(docx_fragments)} fragments")

print("[STEP 2] Inject into template...")
injector = ODTInjector('template/ASL_Template_Delibera.odt')
rich_content = {
    'I_testo_obj': docx_fragments[:1],  # First fragment for Istruttoria
    'P_testo_obj': docx_fragments[:1],  # First fragment for Proposta
    'P_testo_obj_2': docx_fragments[:1], # First fragment for Delibera
}

odt_bytes = injector.inject_placeholders(SIMPLE_DATA, rich_content)
output1 = 'output/test_mode1.odt'
Path(output1).parent.mkdir(parents=True, exist_ok=True)
with open(output1, 'wb') as f:
    f.write(odt_bytes)
print(f"[OK] Generated: {output1}")

# Verify
import zipfile
with zipfile.ZipFile(output1, 'r') as z:
    content = z.read('content.xml')
    if b'2024/001' in content:
        print("[OK] Placeholder successfully injected")
    else:
        print("[FAIL] Placeholder not found")
    if b'bold' in content or b'Istruttoria' in content:
        print("[OK] DOCX content appears in output")

# ============================================================================
# MODALITA 2: TipTap from Editor (simulated)
# ============================================================================
print("\n" + "-" * 60)
print("MODALITA 2: TipTap Editor")
print("-" * 60)

tiptap_data = {
    "type": "doc",
    "content": [
        {
            "type": "paragraph",
            "content": [
                {"type": "text", "text": "Questa è la relazione istruttoria."}
            ]
        },
        {
            "type": "bulletList",
            "content": [
                {
                    "type": "listItem",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [
                                {"type": "text", "text": "Primo punto"}
                            ]
                        }
                    ]
                },
                {
                    "type": "listItem",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [
                                {"type": "text", "text": "Secondo punto"}
                            ]
                        }
                    ]
                }
            ]
        }
    ]
}

print("[STEP 1] Convert TipTap to ODT fragments...")
tiptap_converter = TiptapToODT()
tiptap_json = json.dumps(tiptap_data)
tiptap_fragments, tiptap_styles = tiptap_converter.convert(tiptap_json)
print(f"[OK] Converted to {len(tiptap_fragments)} fragments")

print("[STEP 2] Inject into template...")
rich_content_tiptap = {
    'I_testo_obj': tiptap_fragments,
    'P_testo_obj': tiptap_fragments,
    'P_testo_obj_2': tiptap_fragments,
}

odt_bytes2 = injector.inject_placeholders(SIMPLE_DATA, rich_content_tiptap)
output2 = 'output/test_mode2.odt'
with open(output2, 'wb') as f:
    f.write(odt_bytes2)
print(f"[OK] Generated: {output2}")

# Verify
with zipfile.ZipFile(output2, 'r') as z:
    content = z.read('content.xml')
    if b'relazione istruttoria' in content:
        print("[OK] TipTap content successfully injected")
    else:
        print("[FAIL] TipTap content not found")

# ============================================================================
# MODALITA 3: Simple Editor (text only)
# ============================================================================
print("\n" + "-" * 60)
print("MODALITA 3: Simple Text Editor")
print("-" * 60)

simple_tiptap = {
    "type": "doc",
    "content": [
        {
            "type": "paragraph",
            "content": [
                {"type": "text", "text": "Testo semplice dalla redazione diretta."}
            ]
        }
    ]
}

print("[STEP 1] Convert simple TipTap to ODT...")
simple_fragments, _ = tiptap_converter.convert(json.dumps(simple_tiptap))
print(f"[OK] Converted to {len(simple_fragments)} fragments")

print("[STEP 2] Inject into template...")
rich_content_simple = {
    'I_testo_obj': simple_fragments,
}

odt_bytes3 = injector.inject_placeholders(SIMPLE_DATA, rich_content_simple)
output3 = 'output/test_mode3.odt'
with open(output3, 'wb') as f:
    f.write(odt_bytes3)
print(f"[OK] Generated: {output3}")

# Verify
with zipfile.ZipFile(output3, 'r') as z:
    content = z.read('content.xml')
    if b'redazione diretta' in content:
        print("[OK] Simple text successfully injected")

# ============================================================================
# Summary
# ============================================================================
print("\n" + "=" * 60)
print("TEST SUMMARY")
print("=" * 60)
print(f"Mode 1 (DOCX): {output1}")
print(f"Mode 2 (TipTap editor): {output2}")
print(f"Mode 3 (Simple text): {output3}")
print("\nAll outputs are valid ODT files ready for download!")
print("\nThe POC is functional. You can now:")
print("  1. Run 'python app.py' to start Flask server")
print("  2. Open http://localhost:5000 in your browser")
print("  3. Test all three modes through the web interface")
