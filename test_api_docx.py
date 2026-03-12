#!/usr/bin/env python3
"""Test API endpoint DOCX."""
import sys
import json
import requests
import subprocess
import time
from docx import Document
from io import BytesIO

print("=" * 60)
print("TEST - API DOCX Endpoint")
print("=" * 60)

# Avvia app in background
print("\n[STEP 1] Starting Flask app...")
proc = subprocess.Popen([sys.executable, 'app_simple.py'],
                       stdout=subprocess.DEVNULL,
                       stderr=subprocess.DEVNULL)
time.sleep(2)

try:
    # Test data
    tiptap_data = {
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "IL DIRETTORE GENERALE"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Visto il piano operativo 2024;"}]},
            {"type": "paragraph", "content": [
                {"type": "text", "text": "DELIBERA", "marks": [{"type": "bold"}]}
            ]},
            {"type": "paragraph", "content": [
                {"type": "text", "text": "Di approvare il piano con "},
                {"type": "text", "text": "sottolineato", "marks": [{"type": "underline"}]}
            ]}
        ]
    }

    print("[OK] App started")

    print("\n[STEP 2] POST to /api/genera with editor content...")
    response = requests.post(
        'http://localhost:5001/api/genera',
        data={'editor_content': json.dumps(tiptap_data)}
    )

    if response.status_code != 200:
        print(f"[ERROR] Status {response.status_code}")
        print(response.text)
        sys.exit(1)

    print("[OK] Got response")

    # Verifica che sia un DOCX valido
    docx_bytes = response.content
    print(f"[OK] Received {len(docx_bytes)} bytes")

    # Carica il documento
    doc = Document(BytesIO(docx_bytes))
    print(f"[OK] Valid DOCX with {len(doc.paragraphs)} paragraphs")

    # Verifica il contenuto
    print("\n[STEP 3] Verify content and formatting...")

    text_found = []
    bold_found = False
    underline_found = False

    for para in doc.paragraphs:
        if para.text.strip():
            text_found.append(para.text)
            for run in para.runs:
                if run.bold:
                    bold_found = True
                    print(f"[OK] Bold: '{run.text}'")
                if run.underline:
                    underline_found = True
                    print(f"[OK] Underline: '{run.text}'")

    # Check for required content
    checks = [
        ('DELIBERA' in ' '.join(text_found), "Contains DELIBERA"),
        ('IL DIRETTORE GENERALE' in ' '.join(text_found), "Contains IL DIRETTORE GENERALE"),
        (bold_found, "Contains bold formatting"),
        (underline_found, "Contains underline formatting"),
    ]

    print("\n[STEP 4] Final checks...")
    for check, desc in checks:
        status = "[OK]" if check else "[FAIL]"
        print(f"{status} {desc}")

    if all(c[0] for c in checks):
        print("\n" + "=" * 60)
        print("All tests passed!")
        print("=" * 60)
    else:
        sys.exit(1)

finally:
    proc.terminate()
    proc.wait(timeout=2)
