#!/usr/bin/env python3
"""Crea template DOCX minimalista con placeholder P_testo_obj."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_docx_template():
    output_path = 'template/ASL_Template_Simple.docx'
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Titolo
    title = doc.add_heading('DELIBERA', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Placeholder
    placeholder_para = doc.add_paragraph('P_testo_obj')
    placeholder_para.style = 'Normal'

    doc.save(output_path)
    print(f"[OK] Template creato: {output_path}")

if __name__ == '__main__':
    create_docx_template()
