#!/usr/bin/env python3
"""Test della generazione DOCX."""
import sys
import json
sys.path.insert(0, 'modules')

from tiptap_to_docx import TiptapToDocx
from docx_injector import DocxInjector

print("=" * 60)
print("TEST - DOCX Generation")
print("=" * 60)

# Test con formattazione
tiptap_data = {
    "type": "doc",
    "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "text": "IL DIRETTORE GENERALE"}
        ]},
        {"type": "paragraph", "content": [
            {"type": "text", "text": "Visto il piano operativo 2024;"}
        ]},
        {"type": "paragraph", "content": [
            {"type": "text", "text": "DELIBERA", "marks": [{"type": "bold"}]},
            {"type": "text", "text": " in data"}
        ]},
        {"type": "paragraph", "content": [
            {"type": "text", "text": "Di approvare il piano con "},
            {"type": "text", "text": "ENFASI", "marks": [{"type": "bold"}, {"type": "italic"}]},
            {"type": "text", "text": "."}
        ]}
    ]
}

print("\n[STEP 1] Convert TipTap to DOCX paragraphs...")
converter = TiptapToDocx()
paragraphs, _ = converter.convert(json.dumps(tiptap_data))
print(f"[OK] {len(paragraphs)} paragraphs")

for i, para in enumerate(paragraphs, 1):
    runs_desc = ", ".join([f"{r['text']}" for r in para['runs']])
    print(f"  {i}. {runs_desc}")

print("\n[STEP 2] Inject into template...")
injector = DocxInjector('template/ASL_Template_Simple.docx')
doc = injector.inject_content(paragraphs, 'P_testo_obj')

print("\n[STEP 3] Save to file...")
output_file = 'output/test_docx.docx'
doc.save(output_file)
print(f"[OK] Saved to {output_file}")

print("\n[STEP 4] Verify content...")
from docx import Document
verify_doc = Document(output_file)
print(f"[OK] Document has {len(verify_doc.paragraphs)} paragraphs")

# Check for formatting
has_bold = False
has_italic = False
for para in verify_doc.paragraphs:
    for run in para.runs:
        if run.bold:
            has_bold = True
            print(f"  [OK] Found bold text: '{run.text}'")
        if run.italic:
            has_italic = True
            print(f"  [OK] Found italic text: '{run.text}'")

if not has_bold:
    print("  [CHECK] No bold formatting found")
if not has_italic:
    print("  [CHECK] No italic formatting found")

print("\n" + "=" * 60)
print("Test completed!")
print("=" * 60)
