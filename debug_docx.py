#!/usr/bin/env python3
"""Debug: verify DOCX formatting extraction."""
import sys
import tempfile
sys.path.insert(0, 'modules')

from docx import Document
from docx.shared import Pt

# Create test DOCX with clear formatting
doc = Document()
p = doc.add_paragraph()
p.add_run("Normal ").font.bold = False
p.add_run("BOLD").font.bold = True
p.add_run(" ITALIC").italic = True
p.add_run(" UNDERLINE").underline = True

temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
doc.save(temp_docx)

print(f"Created test DOCX: {temp_docx}")
print("\n--- Analyzing DOCX structure ---")

# Re-open and analyze
doc = Document(temp_docx)
for para in doc.paragraphs:
    print(f"\nParagraph with {len(para.runs)} runs:")
    for i, run in enumerate(para.runs):
        print(f"  Run {i}: '{run.text}'")
        print(f"    - bold: {run.bold}")
        print(f"    - italic: {run.italic}")
        print(f"    - underline: {run.underline}")

# Now test our converter
from docx_to_odt import DocxToODT

print("\n--- Testing DocxToODT conversion ---")
converter = DocxToODT()
fragments, styles = converter.convert_file(temp_docx)

print(f"Fragments generated: {len(fragments)}")
print(f"\nFragment 1:\n{fragments[0][:500]}...")

print(f"\nStyles generated:\n{styles}")

import os
os.unlink(temp_docx)
