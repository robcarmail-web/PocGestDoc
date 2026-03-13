import os
from docx import Document
from modules.docx_injector import DocxInjector

def create_fragment():
    doc = Document()
    doc.add_heading('Fragmento Test', level=3)
    p = doc.add_paragraph('Questo è un paragrafo con ')
    p.add_run('grassetto').bold = True
    p.add_run(' e ')
    p.add_run('corsivo').italic = True
    doc.add_paragraph('Lista: ', style='List Bullet')
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0,0).text = 'Cella 1'
    table.cell(0,1).text = 'Cella 2'
    
    os.makedirs('tmp', exist_ok=True)
    path = 'tmp/test_fragment.docx'
    doc.save(path)
    return path

def test():
    print("1. Creating test fragment...")
    frag_path = create_fragment()
    
    print("2. Running DocxInjector...")
    injector = DocxInjector('template/ASL_Template_Delibera.docx')
    
    simple = {
        'numeroproposta': '2026/999',
        'dataproposta': '13/03/2026',
        'oggetto': 'TEST INIEZIONE DOCX',
        'EstensoreNome': 'Test User',
        'DirGenNome': 'Test Director'
    }
    
    rich = {
        'I_testo_obj': frag_path,
        'P_testo_obj': frag_path,
        'P_testo_obj_2': frag_path
    }
    
    out_bytes = injector.inject_placeholders(simple, rich)
    
    out_path = 'output/test_injection_result.docx'
    with open(out_path, 'wb') as f:
        f.write(out_bytes)
        
    print(f"3. Done! Result saved to {out_path}")

if __name__ == '__main__':
    test()
